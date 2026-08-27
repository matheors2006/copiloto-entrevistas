"""Generates technical summary reports in PDF and Word formats from Markdown-like plain text."""
import os
import re

from docx import Document
from fpdf import FPDF

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _iter_inline_runs(line):
    """Split a line into (text, is_bold) runs based on **bold** markers."""
    position = 0
    for match in BOLD_PATTERN.finditer(line):
        if match.start() > position:
            yield line[position:match.start()], False
        yield match.group(1), True
        position = match.end()
    if position < len(line):
        yield line[position:], False


def _to_latin1(text):
    """Sanitize text for the PDF core font, which only supports the latin-1 charset."""
    return text.encode("latin-1", "replace").decode("latin-1")


def _build_docx(title, content, output_path):
    """Render the report as a .docx file."""
    document = Document()
    document.add_heading(title, level=0)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=2)
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
            continue

        is_bullet = line.startswith("- ")
        text = line[2:] if is_bullet else line
        paragraph = document.add_paragraph(style="List Bullet" if is_bullet else None)
        for run_text, is_bold in _iter_inline_runs(text):
            run = paragraph.add_run(run_text)
            run.bold = is_bold

    document.save(output_path)


def _build_pdf(title, content, output_path):
    """Render the report as a .pdf file."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.multi_cell(0, 10, _to_latin1(title))
    pdf.ln(4)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line:
            pdf.ln(4)
            continue
        if line.startswith("## "):
            pdf.set_font("Helvetica", style="B", size=13)
            pdf.multi_cell(0, 9, _to_latin1(line[3:]))
            continue
        if line.startswith("# "):
            pdf.set_font("Helvetica", style="B", size=15)
            pdf.multi_cell(0, 9, _to_latin1(line[2:]))
            continue

        is_bullet = line.startswith("- ")
        text = line[2:] if is_bullet else line
        pdf.set_font("Helvetica", size=11)
        if is_bullet:
            pdf.write(7, "- ")
        for run_text, is_bold in _iter_inline_runs(text):
            pdf.set_font("Helvetica", style="B" if is_bold else "", size=11)
            pdf.write(7, _to_latin1(run_text))
        pdf.ln(9)

    pdf.output(output_path)


def generate_report(title, content, output_dir="."):
    """Export a technical summary as Reporte_Copiloto.pdf and Reporte_Copiloto.docx in output_dir."""
    os.makedirs(output_dir, exist_ok=True)
    pdf_path = os.path.join(output_dir, "Reporte_Copiloto.pdf")
    docx_path = os.path.join(output_dir, "Reporte_Copiloto.docx")

    _build_docx(title, content, docx_path)
    _build_pdf(title, content, pdf_path)

    return pdf_path, docx_path


REPORT_TITLE = "Reporte de Arquitectura Backend - Copiloto de Entrevistas"

REPORT_CONTENT = """
# Resumen ejecutivo

Este reporte documenta el estado del backend del Copiloto de Entrevistas al cierre de las Fases 1 a 4: captura de audio del sistema, transcripcion en tiempo real, motor de razonamiento LLM, y el servidor que orquesta todo el flujo.

# Fase 1: Captura de audio del sistema (soundcard y Loopback)

- **soundcard** es la libreria que permite acceder a los dispositivos de audio del sistema operativo desde Python, tanto para grabar microfonos normales como para hacer **loopback**.
- El **loopback** es la tecnica que permite capturar el audio que el sistema esta reproduciendo hacia los parlantes o audifonos (por ejemplo, la voz del entrevistador en una llamada de Zoom o Meet), en lugar de depender de un microfono fisico. Esto es posible gracias a **WASAPI** (Windows Audio Session API), que expone el dispositivo de salida por defecto como si fuera un microfono de solo lectura.
- Se eligio capturar en **bloques pequenos (1024 frames)** en lugar de grabaciones largas de una sola vez, porque los bloques grandes generaban advertencias de **data discontinuity** (perdida de continuidad del buffer). Con bloques pequenos, el audio fluye de forma constante y de baja latencia, sin cortes.
- Cada bloque de audio llega como un arreglo de numeros en punto flotante (**float32**, rango -1.0 a 1.0). Antes de enviarlo a cualquier servicio externo se recorta (**clip**) a ese rango y se convierte a **PCM de 16 bits (int16)**, que es el formato que Deepgram espera recibir.

# Fase 2: Transcripcion en tiempo real con Deepgram (WebSockets)

- El audio no se envia como un archivo completo al final: se transmite **en vivo, bloque por bloque, a traves de un WebSocket** hacia los servidores de Deepgram. Un WebSocket mantiene una conexion abierta y bidireccional, ideal para audio continuo, a diferencia de una peticion HTTP tradicional que tendria que esperar a que termine la grabacion.
- Se usa el modelo **nova-2** de Deepgram, configurado con **language="es"**, **encoding="linear16"**, **sample_rate=16000** y **channels=1**, exactamente los mismos parametros con los que se captura y convierte el audio en la Fase 1.
- El parametro **endpointing** (configurado en 900 ms) le indica a Deepgram cuanto silencio debe detectar despues de que alguien habla para considerar que una frase termino. Cuando eso ocurre, Deepgram marca el mensaje como **is_final=True**, que es la senal que el sistema usa para saber que una pregunta del entrevistador esta completa y lista para procesarse.
- Como la captura de audio (soundcard) es una operacion bloqueante y el resto del sistema es asincrono (asyncio), cada bloque se graba dentro de un hilo separado usando **asyncio.to_thread**. Esto evita que el **event loop** se congele mientras se espera el audio, permitiendo que la recepcion de transcripciones ocurra en paralelo sin bloqueos.

# Fase 3: Motor de razonamiento LLM con Groq

- **Groq** es el proveedor que ejecuta el modelo de lenguaje (LLM) a muy baja latencia. El motor implementado en **llm_engine.py** define la funcion asincrona **stream_respuesta(pregunta, cv_json)**, que envia la pregunta transcrita junto con el CV del candidato (en formato JSON) a Groq.
- El modelo configurado actualmente es **openai/gpt-oss-120b**. Originalmente se planeo usar **llama-3.3-70b-versatile**, pero al probarlo en vivo la API respondio con un error 404 (modelo no disponible para esta cuenta), por lo que se verificaron los modelos realmente disponibles y se selecciono el de mayor capacidad entre ellos.
- El **System Prompt** obliga al modelo a responder como un copiloto de entrevistas: maximo 2 o 3 viñetas muy breves (**Talking Points**), sin saludos, basadas estrictamente en el CV, con metricas y tecnologias resaltadas en negrita.
- Se agrego una **regla de contingencia**: si la pregunta es sobre una tecnologia que no aparece en el CV, el modelo debe dar una explicacion tecnica general y precisa, o adaptar la respuesta usando metodologias similares del CV, y tiene prohibido decir frases como "no esta en el CV" o "no lo se". Esto evita que el copiloto deje al candidato sin respuesta en medio de una entrevista real.
- La respuesta se genera con **stream=True**, es decir, Groq envia el texto **token por token** a medida que se genera, en lugar de esperar la respuesta completa. Esto permite mostrar la respuesta en pantalla casi de inmediato.

# Fase 4: FastAPI como orquestador central

- **main.py** expone un servidor **FastAPI** con un unico endpoint WebSocket en **/ws**. Cuando un cliente (la futura app de escritorio) se conecta, el servidor acepta la conexion y arranca todo el pipeline para esa sesion.
- Al aceptar la conexion, el servidor: (1) carga el CV actualizado desde disco, (2) abre su propia conexion WebSocket hacia Deepgram con los parametros de la Fase 2, y (3) lanza en paralelo la captura de audio del sistema y el envio continuo hacia Deepgram.
- Cuando Deepgram devuelve una transcripcion final (**is_final=True**), FastAPI hace dos cosas en secuencia: primero envia al cliente un mensaje **{"type": "question", "text": "..."}** con la pregunta detectada, y de inmediato llama a **stream_respuesta** del motor LLM, reenviando cada token que llega de Groq al cliente como **{"type": "answer_token", "text": "..."}**.
- FastAPI actua asi como el **orquestador**: coordina tres sistemas externos (el microfono via soundcard, Deepgram para voz-a-texto, y Groq para el razonamiento) sin que ninguno bloquee a los demas, usando tareas asincronas (**asyncio.create_task**) que corren de forma concurrente.
- Cuando el cliente se desconecta, el servidor detecta el cierre mediante una tarea que escucha la desconexion (**websocket.receive_text**), y cancela de forma ordenada las tareas de audio y de Deepgram que quedaban en segundo plano, evitando procesos huerfanos o conexiones abiertas innecesariamente.

# Conclusion

Con estas cuatro fases, el backend ya cubre el flujo completo de un copiloto de entrevistas en tiempo real: escuchar el audio del sistema, convertirlo en texto, generar sugerencias de respuesta basadas en el CV del candidato, y transmitir todo de vuelta a un cliente conectado, listo para integrarse con el frontend de escritorio en Tauri.
"""


if __name__ == "__main__":
    pdf_path, docx_path = generate_report(REPORT_TITLE, REPORT_CONTENT)
    print(f"PDF generado en: {pdf_path}")
    print(f"Word generado en: {docx_path}")
